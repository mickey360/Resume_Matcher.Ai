from io import BytesIO

import fitz
from docx import Document


def extract_text(data: bytes, filename: str) -> str:
    name = filename.lower()

    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)

    raise ValueError("Unsupported resume format. Use PDF or DOCX.")


def _extract_pdf(data: bytes) -> str:
    with fitz.open(stream=data, filetype="pdf") as document:
        pages = [page.get_text("text") for page in document]
    return "\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(parts).strip()
